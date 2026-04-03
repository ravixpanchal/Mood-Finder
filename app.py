import streamlit as st
import librosa
import numpy as np
import io
import os
import tempfile
from datetime import datetime

from utils import plot_waveform, plot_spectrogram, plot_mfcc, extract_features
from emotion_engine import classify_emotion

# ── Theme (Dark — fixed) ────────────────────────────────────────────────────
BG          = "#0b0914"
GRID_COLOR  = "rgba(255,255,255,0.03)"
TEXT_COLOR  = "#ffffff"
SUBTLE      = "#94a3b8"
CARD_BG     = "rgba(255,255,255,0.04)"
CARD_BORDER = "rgba(255,255,255,0.08)"
BADGE_BG    = "rgba(255,255,255,0.03)"
BADGE_BOR   = "rgba(255,255,255,0.10)"
BADGE_TXT   = "#cbd5e1"
UPLOAD_BG   = "rgba(99,102,241,0.04)"
UPLOAD_BOR  = "rgba(99,102,241,0.25)"
FOOTER_TXT  = "#64748b"
SUCCESS_BG  = "rgba(34,197,94,0.10)"
SUCCESS_BOR = "rgba(34,197,94,0.30)"
SUCCESS_TXT = "#4ade80"

# ── Page Config ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Emotion Detection from Speech",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ── Global CSS ───────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

*, *::before, *::after {{ box-sizing: border-box; }}

html, body, .stApp {{
    font-family: "Inter", sans-serif;
    background-color: {BG} !important;
    background-image:
        linear-gradient({GRID_COLOR} 1px, transparent 1px),
        linear-gradient(90deg, {GRID_COLOR} 1px, transparent 1px);
    background-size: 40px 40px;
    color: {TEXT_COLOR} !important;
    transition: background-color 0.3s ease, color 0.3s ease;
}}

.block-container {{
    padding-top: 1.5rem !important;
    padding-left: clamp(0.75rem, 3vw, 2.5rem) !important;
    padding-right: clamp(0.75rem, 3vw, 2.5rem) !important;
    max-width: 100% !important;
    width: 100% !important;
}}

/* Force Streamlit root containers to fill viewport */
.main > div, [data-testid="stAppViewContainer"], [data-testid="stMain"] {{
    max-width: 100% !important;
    width: 100% !important;
    padding: 0 !important;
}}
section[data-testid="stSidebar"] + div {{
    max-width: 100% !important;
}}

/* ── Hero ── */
.hero-container {{ text-align: center; padding: 20px 0 35px; }}
.top-badge {{
    display: inline-block; padding: 6px 16px; border-radius: 50px;
    background: {SUCCESS_BG}; border: 1px solid {SUCCESS_BOR};
    color: {SUCCESS_TXT}; font-size: 11px; font-weight: 600;
    letter-spacing: 2px; margin-bottom: 22px;
}}
.main-title {{
    font-size: clamp(2rem, 6vw, 4rem);
    font-weight: 800; line-height: 1.1;
    margin-bottom: 18px; color: {TEXT_COLOR};
}}
.highlight {{ color: #00d2ff; }}
.subtitle {{
    color: {SUBTLE}; font-size: clamp(0.88rem, 2vw, 1.05rem);
    max-width: 620px; margin: 0 auto; line-height: 1.65;
}}
.badges-row {{
    display: flex; justify-content: center; flex-wrap: wrap;
    gap: 10px; margin-top: 32px;
}}
.feature-badge {{
    padding: 7px 14px; border-radius: 50px;
    background: {BADGE_BG}; border: 1px solid {BADGE_BOR};
    color: {BADGE_TXT}; font-size: 12px;
    display: flex; align-items: center; gap: 8px;
}}


/* ── Metrics ── */
[data-testid="stMetric"] {{
    background: {CARD_BG};
    border: 1px solid {CARD_BORDER};
    border-radius: 12px;
    padding: 12px 16px;
    margin-bottom: 8px;
}}
[data-testid="stMetricLabel"] p {{ color: {SUBTLE} !important; font-size: 12px !important; }}
[data-testid="stMetricValue"] {{ color: {TEXT_COLOR} !important; }}

/* ── File uploader ── */
[data-testid="stFileUploader"] {{
    background: {UPLOAD_BG};
    border: 2px dashed {UPLOAD_BOR};
    border-radius: 16px;
    padding: 12px;
}}

/* ── Progress bar ── */
[data-testid="stProgress"] > div > div {{
    background: linear-gradient(90deg, #6366f1, #22d3ee);
}}

/* ── Section headers ── */
h1, h2, h3, h4 {{ color: {TEXT_COLOR} !important; }}

/* ── Download buttons ── */
.dl-row {{
    display: flex; gap: 12px; flex-wrap: wrap; margin-top: 20px;
}}
.dl-card {{
    flex: 1; min-width: 160px;
    background: {CARD_BG}; border: 1px solid {CARD_BORDER};
    border-radius: 14px; padding: 16px 20px;
    text-align: center;
}}

/* ── Footer ── */
.site-footer {{
    text-align: center;
    color: {FOOTER_TXT};
    font-size: 13px;
    padding: 30px 0 10px;
    border-top: 1px solid {CARD_BORDER};
    margin-top: 40px;
}}

/* ── Responsive ── */
@media (max-width: 900px) {{
    .main-title {{ font-size: 2.2rem; }}
    .subtitle   {{ font-size: 0.92rem; padding: 0 12px; }}
    .badges-row {{ gap: 8px; }}
}}
@media (max-width: 640px) {{
    .main-title  {{ font-size: 1.7rem; }}
    .subtitle    {{ font-size: 0.85rem; padding: 0 8px; }}
    .feature-badge {{ font-size: 11px; padding: 5px 10px; }}
    [data-testid="column"] {{ min-width: 100% !important; flex: 1 1 100% !important; }}
    .dl-row {{ flex-direction: column; }}
    [data-testid="stMetric"] {{ padding: 10px 12px; }}
}}

/* ── Streamlit markdown text colour ── */
p, li, label, .stMarkdown {{ color: {TEXT_COLOR} !important; }}

/* ── Success box ── */
[data-testid="stAlert"] {{
    background: {SUCCESS_BG} !important;
    border: 1px solid {SUCCESS_BOR} !important;
    color: {SUCCESS_TXT} !important;
}}
</style>
""", unsafe_allow_html=True)


# ── Hero ─────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="hero-container">
  <div class="top-badge">🟢 EMOTION AI • SPEECH ANALYSIS</div>
  <div class="main-title">Hear what words<br><span class="highlight">can't always say</span></div>
  <div class="subtitle">Upload speech audio and let our system decode the emotional state hidden within
  acoustic patterns — with real-time visualizations.</div>
  <div class="badges-row">
    <div class="feature-badge">🌊 Waveform Analysis</div>
    <div class="feature-badge">📊 Spectrogram</div>
    <div class="feature-badge">🎨 MFCC</div>
    <div class="feature-badge">🤖 Feature-based Classification</div>
  </div>
</div>
""", unsafe_allow_html=True)

st.markdown("---")


# ── Audio Loader ──────────────────────────────────────────────────────────────
def _get_ffmpeg_path():
    """Return path to a working ffmpeg binary (system PATH or imageio-ffmpeg bundle)."""
    import shutil
    if shutil.which("ffmpeg"):
        return shutil.which("ffmpeg")
    try:
        import imageio_ffmpeg
        path = imageio_ffmpeg.get_ffmpeg_exe()
        if path and os.path.isfile(path):
            return path
    except Exception:
        pass
    return None


def _decode_with_ffmpeg(file_bytes: bytes, ext: str, ffmpeg_path: str) -> bytes:
    """
    Use ffmpeg directly (subprocess) to convert any audio format → WAV bytes.
    This completely bypasses pydub's ffprobe dependency.
    """
    import subprocess
    cmd = [
        ffmpeg_path,
        "-y",                  # overwrite output
        "-hide_banner",
        "-loglevel", "error",
        "-i", "pipe:0",        # read from stdin
        "-f", "wav",           # output format
        "-ar", "22050",        # resample to 22050 Hz (librosa default)
        "-ac", "1",            # mono
        "pipe:1",              # write to stdout
    ]
    result = subprocess.run(
        cmd,
        input=file_bytes,
        capture_output=True,
        timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"ffmpeg conversion failed: {result.stderr.decode(errors='replace')}"
        )
    return result.stdout


def load_audio_bytes(file_bytes: bytes, ext: str):
    """
    Load audio from raw bytes using a multi-strategy approach:
      1. WAV / FLAC / OGG / AIFF  → librosa/soundfile directly (no ffmpeg needed)
      2. Everything else           → ffmpeg subprocess via stdin/stdout pipe
                                     (no ffprobe, no pydub dependency issues)
      3. Last resort               → temp file + librosa
    """
    ext = ext.lower()

    # ── Strategy 1: native librosa/soundfile formats ─────────────────────────
    if ext in ("wav", "flac", "ogg", "aiff", "aif"):
        try:
            y, sr = librosa.load(io.BytesIO(file_bytes), sr=None, mono=True)
            return y, sr
        except Exception:
            pass  # fall through

    # ── Strategy 2: ffmpeg subprocess (pipe) — no ffprobe needed ─────────────
    ffmpeg_path = _get_ffmpeg_path()
    if ffmpeg_path:
        try:
            wav_bytes = _decode_with_ffmpeg(file_bytes, ext, ffmpeg_path)
            y, sr = librosa.load(io.BytesIO(wav_bytes), sr=None, mono=True)
            return y, sr
        except Exception as ffmpeg_err:
            pass  # fall through to temp-file strategy

    # ── Strategy 3: temp file + librosa ──────────────────────────────────────
    try:
        suffix = f".{ext}"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            y, sr = librosa.load(tmp_path, sr=None, mono=True)
            return y, sr
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    except Exception as tmp_err:
        pass

    # ── All strategies failed — raise clear error ─────────────────────────────
    if not ffmpeg_path:
        raise RuntimeError(
            f"Cannot load '{ext.upper()}' — ffmpeg not found.\n"
            "Install it with:  pip install imageio-ffmpeg\n"
            "Or system-wide:   https://ffmpeg.org/download.html"
        )
    raise RuntimeError(
        f"Cannot decode '{ext.upper()}' audio even with ffmpeg. "
        "Try converting the file to WAV first."
    )


# ── Export Helpers ─────────────────────────────────────────────────────────────
def build_pdf_bytes(emotion, confidence, scores, features, filename):
    """Generate a simple PDF report using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch
    except ImportError:
        return None

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.9*inch, rightMargin=0.9*inch,
                            topMargin=0.9*inch, bottomMargin=0.9*inch)
    styles = getSampleStyleSheet()
    story  = []

    title_style = ParagraphStyle("title", parent=styles["Title"],
                                 fontSize=22, textColor=colors.HexColor("#0f172a"),
                                 spaceAfter=6)
    h2_style    = ParagraphStyle("h2", parent=styles["Heading2"],
                                 fontSize=14, textColor=colors.HexColor("#1e40af"),
                                 spaceBefore=14, spaceAfter=6)
    body_style  = ParagraphStyle("body", parent=styles["Normal"],
                                 fontSize=11, textColor=colors.HexColor("#334155"),
                                 leading=16)

    story.append(Paragraph("🎙️ Emotion Detection Report", title_style))
    story.append(Paragraph(f"File: {filename}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
    story.append(Spacer(1, 14))

    story.append(Paragraph("Detected Emotion", h2_style))
    story.append(Paragraph(f"<b>{emotion}</b> — Confidence: {confidence:.1f}%", body_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("All Emotion Scores", h2_style))
    table_data  = [["Emotion", "Score (%)"]]
    for emo, sc in sorted(scores.items(), key=lambda x: -x[1]):
        table_data.append([emo, f"{sc:.1f}%"])

    tbl = Table(table_data, colWidths=[3*inch, 2*inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0),  colors.HexColor("#1e40af")),
        ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 11),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
        ("GRID",        (0, 0), (-1, -1), 0.4, colors.HexColor("#e2e8f0")),
        ("PADDING",     (0, 0), (-1, -1), 8),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 14))

    story.append(Paragraph("Acoustic Features", h2_style))
    feat_data = [["Feature", "Value"],
                 ["Zero Crossing Rate",  f"{features['zcr']:.4f}"],
                 ["RMS Energy",          f"{features['rms']:.4f}"],
                 ["Spectral Centroid",   f"{features['centroid']:.1f} Hz"],
                 ["Spectral Rolloff",    f"{features['rolloff']:.1f} Hz"],
                 ["Tempo",               f"{features['tempo']:.1f} BPM"]]
    ftbl = Table(feat_data, colWidths=[3*inch, 2*inch])
    ftbl.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0),  colors.HexColor("#0f766e")),
        ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 11),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0fdf4"), colors.white]),
        ("GRID",        (0, 0), (-1, -1), 0.4, colors.HexColor("#d1fae5")),
        ("PADDING",     (0, 0), (-1, -1), 8),
    ]))
    story.append(ftbl)
    story.append(Spacer(1, 20))
    story.append(Paragraph("Made with ❤️ by Team | All Rights Reserved © 2026", body_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def build_excel_bytes(emotion, confidence, scores, features, filename):
    """Generate an Excel report using openpyxl."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        return None

    wb = openpyxl.Workbook()

    # ── Sheet 1: Summary ──────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Summary"
    ws1.column_dimensions["A"].width = 28
    ws1.column_dimensions["B"].width = 22

    def hdr(ws, row, text, bg="1e40af"):
        c = ws.cell(row=row, column=1, value=text)
        c.font      = Font(bold=True, color="FFFFFF", size=12)
        c.fill      = PatternFill("solid", fgColor=bg)
        c.alignment = Alignment(horizontal="center")
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)

    hdr(ws1, 1, "🎙️ Emotion Detection Report")
    ws1.cell(2, 1, "File").font       = Font(bold=True)
    ws1.cell(2, 2, filename)
    ws1.cell(3, 1, "Generated").font  = Font(bold=True)
    ws1.cell(3, 2, datetime.now().strftime("%Y-%m-%d %H:%M"))
    ws1.cell(4, 1, "Detected Emotion").font = Font(bold=True)
    ws1.cell(4, 2, emotion)
    ws1.cell(5, 1, "Confidence (%)").font   = Font(bold=True)
    ws1.cell(5, 2, round(confidence, 1))

    hdr(ws1, 7, "All Emotion Scores", bg="0f766e")
    ws1.cell(8, 1, "Emotion").font  = Font(bold=True)
    ws1.cell(8, 2, "Score (%)").font = Font(bold=True)
    for i, (emo, sc) in enumerate(sorted(scores.items(), key=lambda x: -x[1]), start=9):
        ws1.cell(i, 1, emo)
        ws1.cell(i, 2, round(sc, 1))

    # ── Sheet 2: Features ─────────────────────────────────────────────
    ws2 = wb.create_sheet("Acoustic Features")
    ws2.column_dimensions["A"].width = 26
    ws2.column_dimensions["B"].width = 20
    hdr(ws2, 1, "Acoustic Features", bg="7c3aed")
    headers = ["Zero Crossing Rate", "RMS Energy", "Spectral Centroid (Hz)",
               "Spectral Rolloff (Hz)", "Tempo (BPM)"]
    vals    = [round(features["zcr"], 4), round(features["rms"], 4),
               round(features["centroid"], 1), round(features["rolloff"], 1),
               round(features["tempo"], 1)]
    for i, (h, v) in enumerate(zip(headers, vals), start=2):
        ws2.cell(i, 1, h).font = Font(bold=True)
        ws2.cell(i, 2, v)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def build_word_bytes(emotion, confidence, scores, features, filename):
    """Generate a Word (.docx) report using python-docx."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = DocxDocument()
    doc.core_properties.author = "Emotion Detection AI"

    # Title
    title = doc.add_heading("🎙️ Emotion Detection Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"File: {filename}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    doc.add_heading("Detected Emotion", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"{emotion}  —  Confidence: {confidence:.1f}%")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_heading("All Emotion Scores", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Emotion"
    hdr_cells[1].text = "Score (%)"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for emo, sc in sorted(scores.items(), key=lambda x: -x[1]):
        row_cells = table.add_row().cells
        row_cells[0].text = emo
        row_cells[1].text = f"{sc:.1f}%"

    doc.add_paragraph()
    doc.add_heading("Acoustic Features", level=1)
    feat_tbl = doc.add_table(rows=1, cols=2)
    feat_tbl.style = "Table Grid"
    fh = feat_tbl.rows[0].cells
    fh[0].text, fh[1].text = "Feature", "Value"
    for cell in fh:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows = [("Zero Crossing Rate",  f"{features['zcr']:.4f}"),
            ("RMS Energy",          f"{features['rms']:.4f}"),
            ("Spectral Centroid",   f"{features['centroid']:.1f} Hz"),
            ("Spectral Rolloff",    f"{features['rolloff']:.1f} Hz"),
            ("Tempo",               f"{features['tempo']:.1f} BPM")]
    for feat, val in rows:
        rc = feat_tbl.add_row().cells
        rc[0].text, rc[1].text = feat, val

    doc.add_paragraph()
    footer_p = doc.add_paragraph("Made with ❤️ by Team | All Rights Reserved © 2026")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── File Uploader ────────────────────────────────────────────────────────────
st.markdown("### 🎙️ Upload Audio File")

SUPPORTED = ["wav", "mp3", "mp4", "ogg", "flac", "aac", "m4a", "wma", "aiff", "opus"]
uploaded_file = st.file_uploader(
    "Supports WAV · MP3 · MP4 · OGG · FLAC · AAC · M4A · WMA · AIFF · OPUS",
    type=SUPPORTED
)

if uploaded_file is not None:
    ext      = uploaded_file.name.rsplit(".", 1)[-1].lower()
    mime_map = {"wav": "audio/wav", "mp3": "audio/mpeg", "mp4": "video/mp4",
                "ogg": "audio/ogg", "flac": "audio/flac", "aac": "audio/aac",
                "m4a": "audio/mp4", "wma": "audio/x-ms-wma",
                "aiff": "audio/aiff", "opus": "audio/opus"}
    mime = mime_map.get(ext, "audio/wav")
    st.audio(uploaded_file, format=mime)

    try:
        file_bytes = uploaded_file.getvalue()

        with st.spinner(f"Loading {ext.upper()} audio…"):
            y, sr = load_audio_bytes(file_bytes, ext)

        st.success(
            f"✅ Audio loaded!  |  **Format:** {ext.upper()}  |  "
            f"**Sample Rate:** {sr} Hz  |  **Duration:** {len(y)/sr:.1f}s",
            icon="🎵"
        )

        # ── Feature Extraction & Visualizations ──────────────────────────
        st.markdown("---")
        st.header("🔍 Audio Feature Analysis")

        col1, col2 = st.columns([1, 1])
        with col1:
            st.subheader("Waveform")
            st.markdown("How the amplitude of the audio signal changes over time.")
            st.pyplot(plot_waveform(y, sr), use_container_width=True)

        with col2:
            st.subheader("Spectrogram")
            st.markdown("Frequencies of the signal as they vary with time.")
            st.pyplot(plot_spectrogram(y, sr), use_container_width=True)

        st.markdown("---")

        col3, col4 = st.columns([2, 1])
        with col3:
            st.subheader("MFCC (Mel Frequency Cepstral Coefficients)")
            st.markdown("Short-term power spectrum — highly useful for speech processing.")
            st.pyplot(plot_mfcc(y, sr), use_container_width=True)

        with col4:
            st.subheader("Extracted Features")
            features = extract_features(y, sr)
            st.metric("Mean ZCR",          f"{features['zcr']:.4f}")
            st.metric("RMS Energy",        f"{features['rms']:.4f}")
            st.metric("Spectral Centroid", f"{features['centroid']:.1f} Hz")
            st.metric("Spectral Rolloff",  f"{features['rolloff']:.1f} Hz")
            st.metric("Tempo (BPM)",       f"{features['tempo']:.1f}")

        st.markdown("---")

        # ── Emotion Classification ────────────────────────────────────────
        st.header("🧠 Emotion Classification")
        emotion, confidence, scores = classify_emotion(features)

        emoji_map = {
            "Happy":   "😄", "Sad":     "😢", "Angry":   "😡",
            "Fearful": "😨", "Neutral": "😐", "Calm":    "😌",
        }
        emoji = emoji_map.get(emotion, "🤔")

        st.success(f"### Detected Emotion: {emoji} **{emotion}**")
        st.progress(confidence / 100, text=f"Confidence: {confidence:.1f}%")

        st.markdown("#### All Emotion Scores")
        score_col1, score_col2 = st.columns(2)
        for i, (emo, score) in enumerate(sorted(scores.items(), key=lambda x: -x[1])):
            col = score_col1 if i % 2 == 0 else score_col2
            with col:
                st.metric(f"{emoji_map.get(emo, '🤔')} {emo}", f"{score:.1f}%")

        # ── Download Section ──────────────────────────────────────────────
        st.markdown("---")
        st.header("📥 Download Report")
        st.markdown("Export the emotion analysis results in your preferred format.")

        dl_col1, dl_col2, dl_col3 = st.columns(3)
        base_name = uploaded_file.name.rsplit(".", 1)[0]

        with dl_col1:
            pdf_bytes = build_pdf_bytes(emotion, confidence, scores, features, uploaded_file.name)
            if pdf_bytes:
                st.download_button(
                    label="📄 Download PDF",
                    data=pdf_bytes,
                    file_name=f"{base_name}_emotion_report.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            else:
                st.info("Install `reportlab` to enable PDF export:\n`pip install reportlab`")

        with dl_col2:
            xl_bytes = build_excel_bytes(emotion, confidence, scores, features, uploaded_file.name)
            if xl_bytes:
                st.download_button(
                    label="📊 Download Excel",
                    data=xl_bytes,
                    file_name=f"{base_name}_emotion_report.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            else:
                st.info("Install `openpyxl` to enable Excel export:\n`pip install openpyxl`")

        with dl_col3:
            word_bytes = build_word_bytes(emotion, confidence, scores, features, uploaded_file.name)
            if word_bytes:
                st.download_button(
                    label="📝 Download Word",
                    data=word_bytes,
                    file_name=f"{base_name}_emotion_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.info("Install `python-docx` to enable Word export:\n`pip install python-docx`")

    except Exception as e:
        err_msg = str(e)
        st.error(f"❌ {err_msg}")
        if "ffmpeg" in err_msg.lower() or "imageio" in err_msg.lower():
            st.info("💡 Run: `pip install imageio-ffmpeg` then restart Streamlit. WAV files always work without ffmpeg.")
        else:
            st.info("💡 Try converting your file to WAV format and re-uploading.")

# ── Footer ───────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="site-footer">
    Made with ❤️ by Team. All Rights Reserved © 2026
</div>
""", unsafe_allow_html=True)
